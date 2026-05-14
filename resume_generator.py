"""Generate tailored resume files (.docx) for each job application."""
from pathlib import Path
from docx import Document
import re

TAILORED_DIR = Path(__file__).parent / "tailored_resumes"


def save_tailored_resume(tailored_text: str, company: str, job_title: str) -> str:
    """Save tailored resume as .docx and return the file path."""
    TAILORED_DIR.mkdir(exist_ok=True)
    
    # Clean filename
    safe_company = re.sub(r'[^\w\s-]', '', company).strip().replace(' ', '_')
    safe_title = re.sub(r'[^\w\s-]', '', job_title).strip().replace(' ', '_')
    filename = f"Resume_{safe_company}_{safe_title}.docx"
    filepath = TAILORED_DIR / filename
    
    doc = Document()
    for para in tailored_text.split('\n'):
        para = para.strip()
        if not para:
            continue
        # Detect headings (all caps or short lines ending with colon)
        if para.isupper() or (len(para) < 50 and para.endswith(':')):
            doc.add_heading(para.rstrip(':'), level=2)
        else:
            doc.add_paragraph(para)
    
    doc.save(str(filepath))
    return str(filepath)
